import os
from docx import Document
from tkinter import filedialog, Tk

def get_files_from_directory(directory_path):
    return [os.path.join(directory_path, f) for f in os.listdir(directory_path) if f.endswith(('.docx', '.DOCX'))]

def modify_header_and_table(doc_path, header_data, table_data, number):
    doc = Document(doc_path)

    # Изменение заголовка
    for paragraph in doc.paragraphs:
        if "Заявки на потребность №" in paragraph.text and "от " in paragraph.text:
            # Очищаем все run и вставляем новый текст
            for run in paragraph.runs:
                run.text = ""
            paragraph.runs[0].text = f"Заявки на потребность № {number} от {header_data}"

    # Изменение таблицы
    if len(doc.tables) > 0:
        table = doc.tables[0]
        if len(table.rows) >= 4 and len(table.rows[3].cells) >= 3:
            original_text = table.rows[3].cells[2].text.strip()
            
            # Ищем последний пробел перед временем
            last_space_index = original_text.rfind(' ')
            
            if last_space_index != -1:
                # Сохраняем время после последнего пробела
                time_part = original_text[last_space_index+1:]
                # Вставляем новые данные + сохраненное время
                new_text = f"{table_data} {time_part}"
            else:
                # Если пробела нет, заменяем весь текст
                new_text = table_data
                
            table.rows[3].cells[2].text = new_text

    # Сохранение в новую папку
    directory, filename = os.path.split(doc_path)
    new_folder = os.path.join(directory, "Измененные заявки")
    os.makedirs(new_folder, exist_ok=True)

    name, ext = os.path.splitext(filename)
    new_name = f"{name.split('М')[0]}М{number}{ext}"
    new_path = os.path.join(new_folder, new_name)

    doc.save(new_path)

def main():
    root = Tk()
    root.withdraw()
    print("Выберите папку с документами: ")

    folder_path = filedialog.askdirectory(title="Выберите папку с документами")
    
    if not folder_path:
        print("Папка не выбрана.")
        return

    number = input("Введите начальный номер первого файла: ")
    try:
        number = int(number)
    except ValueError:
        print("Ошибка: Введенное значение не является целым числом.")
        return

    header_data = input("Введите новые данные для заголовка: ")
    table_data = input("Введите новые данные для таблицы: ")

    print("")

    doc_files = get_files_from_directory(folder_path)
    
    if (doc_files):
        pos = 1
    else:
        pos = 0

    for doc_file in doc_files:
        print(f"Обрабатывается файл №{pos}: {doc_file}")
        modify_header_and_table(doc_file, header_data, table_data, number)
        pos += 1
        number += 1
    
    print(f"\nОбработано файлов - {pos - 1}.\n")
    input("Нажмите Enter для выхода... ")

if __name__ == "__main__":
    main()

# E:\Programs\Python_3.10.11\python.exe -m PyInstaller --onefile main.py